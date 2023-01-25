import docx
"""How to open word documents, edit style and create new docs"""

# open doc
d = docx.Document('demo.docx')

# print first paragraph
print(d.paragraphs[0].text)


p = d.paragraphs[1]
print(p)

# each style is a run, here we edit run 3 to be underlined and change the text
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined'

# save changes
d.save('demo2.docx')

# check style of paragraph
print(p.style)

# change style of paragraph
p.style = 'Title'

d.save('demo3.docx')

# Create a doc

d = docx.Document()

d.add_paragraph('Hello, this is a paragraph.')
d.add_paragraph('This is another paragraph.')
d.save('demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs[1].bold = True
d.save('demo5.docx')
